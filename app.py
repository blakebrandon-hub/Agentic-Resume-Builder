import json
import os
import re
import threading
import uuid
import logging
from io import BytesIO

from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches, Pt
import PyPDF2

from google import genai
from google.genai import types

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ---------------------------------------------------
# Configuration & Globals
# ---------------------------------------------------

MODEL = "gemini-2.5-flash"  # Using the stable/latest standard flash model
MAX_ITERATIONS_RESUME = 3
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "outputs")
os.makedirs(OUTPUT_DIR, exist_ok=True)

app = Flask(__name__)
CORS(app)

# Configure Gemini Client
api_key = os.environ.get("GEMINI_API_KEY", "YOUR_GEMINI_API_KEY_HERE")
client = genai.Client(api_key=api_key.strip())

# Job Store
resume_jobs = {}

# ---------------------------------------------------
# Core AI Utilities
# ---------------------------------------------------

def call_llm(system_prompt, user_prompt, temperature=0.3):
    """Make an LLM call using the Google GenAI SDK."""
    try:
        response = client.models.generate_content(
            model=MODEL,
            contents=user_prompt,
            config=types.GenerateContentConfig(
                system_instruction=system_prompt,
                temperature=temperature
            )
        )
        
        result = response.text.strip()
        
        # Safely clean up JSON wrapped in markdown
        if "JSON ONLY" in system_prompt and result.startswith("```"):
            result = re.sub(r'^```(?:json)?\n', '', result)
            result = re.sub(r'\n```$', '', result)
            
        return result
    except Exception as e:
        logger.error(f"Gemini API Error: {str(e)}")
        if "JSON ONLY" in system_prompt:
            return "{}"
        return ""

class Agent:
    def __init__(self, name, system_prompt):
        self.name = name
        self.system_prompt = system_prompt

    def run(self, prompt, temperature=0.3):
        return call_llm(self.system_prompt, prompt, temperature)

# ---------------------------------------------------
# File Extraction Utilities
# ---------------------------------------------------

def extract_text_from_file(file):
    """Extract text from uploaded PDF, DOCX, or plain text files."""
    filename = secure_filename(file.filename).lower()
    try:
        if filename.endswith('.pdf'):
            pdf = PyPDF2.PdfReader(file)
            return "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
        elif filename.endswith('.docx'):
            doc = Document(file)
            return "\n".join(para.text for para in doc.paragraphs)
        else:
            return file.read().decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"Failed to extract text from {filename}: {e}")
        return ""

# ---------------------------------------------------
# Resume Builder Agents & Logic
# ---------------------------------------------------

resume_analyzer = Agent("resume_analyzer", """
You are a resume analysis agent. Analyze the job description and the user's target role.
Extract key requirements, technical skills, soft skills, and what experience to emphasize.
Return JSON ONLY:
{
  "key_skills": ["skill1", "skill2"],
  "requirements": ["req1", "req2"],
  "priorities": ["what to emphasize"]
}
""")

resume_writer = Agent("resume_writer", """
You are an elite professional resume writer. Create an ATS-optimized resume.
- Tailor the candidate's existing experience directly to the job description and target role.
- Use strong action verbs and highlight quantifiable achievements.
- Format strictly for easy ATS parsing.
- Keep the format clean: use ### for major section headers (e.g., ### EXPERIENCE, ### SKILLS).
- Do not include conversational filler; output only the resume text.
""")

resume_critic = Agent("resume_critic", """
You are a resume critic. Evaluate the resume's quality, formatting, and ATS compatibility based on the job description.
Return JSON ONLY:
{
  "pass": true or false,
  "feedback": "detailed critique on what is missing or needs improvement"
}
Pass = true ONLY if the resume perfectly addresses the job description and has zero formatting issues.
""")

def create_resume_docx(resume_text, job_id):
    """Create a nicely formatted Word document from the AI generated text."""
    doc = Document()
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    for line in resume_text.split('\n'):
        if not line.strip():
            continue
            
        p = doc.add_paragraph()
        # Handle custom headings formatted as ### HEADING
        if line.startswith('###'):
            p.add_run(line.replace('###', '').strip()).bold = True
            p.runs[0].font.size = Pt(14)
        # Handle bolded text (**text**)
        elif '**' in line:
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Inside ** **
                    p.add_run(part).bold = True
                else:
                    p.add_run(part)
        # Standard bullet points
        elif line.strip().startswith('-') or line.strip().startswith('•'):
            p.style = 'List Bullet'
            p.add_run(line.strip()[1:].strip())
        # Standard upper case headings
        elif line.isupper() and len(line) > 3:
            p.add_run(line).bold = True
            p.runs[0].font.size = Pt(12)
        else:
            p.add_run(line)
            
    output_path = os.path.join(OUTPUT_DIR, f"resume_{job_id}.docx")
    doc.save(output_path)
    return output_path

def run_resume_agent(target_role, job_description, current_resume, job_id):
    state = {
        "target_role": target_role,
        "job_description": job_description,
        "current_resume": current_resume,
        "analysis": {},
        "resume": "",
        "feedback": "",
        "docx_path": None,
        "status": "running",
        "current_phase": "Extracting Context",
        "current_iteration": 0,
        "logs": []
    }
    resume_jobs[job_id] = state

    try:
        for step in range(MAX_ITERATIONS_RESUME):
            state["current_iteration"] = step + 1
            state["logs"].append(f"Starting Iteration {step + 1}/{MAX_ITERATIONS_RESUME}")

            # Analyzer Phase
            state["current_phase"] = "Analyzing Job Profile"
            state["logs"].append("Analyzing job description and target role...")
            try:
                analysis_prompt = f"Target Role: {state['target_role']}\n\nJob Description:\n{state['job_description']}"
                analysis_response = resume_analyzer.run(analysis_prompt, temperature=0.2)
                state["analysis"] = json.loads(analysis_response)
                state["logs"].append("Analysis complete. Key requirements extracted.")
            except Exception as e:
                state["logs"].append(f"Analyzer warning: Could not parse JSON. Proceeding with raw analysis.")

            # Writer Phase
            state["current_phase"] = "Drafting Resume"
            state["logs"].append("Tailoring resume content to match requirements...")
            writer_prompt = f"""
            Target Role: {state['target_role']}
            Job Description:\n{state['job_description']}
            
            Candidate's Current Background/Resume Data:\n{state['current_resume']}
            
            Priorities/Analysis:\n{json.dumps(state['analysis'], indent=2)}
            
            Create a highly tailored resume that highlights relevant experience and skills. 
            Format cleanly using ### for section headers. Incorporate previous feedback if any: {state['feedback']}
            """
            state["resume"] = resume_writer.run(writer_prompt, temperature=0.3)

            # Critic Phase
            state["current_phase"] = "Reviewing Draft"
            state["logs"].append("Reviewing resume for ATS compatibility and impact...")
            try:
                review_prompt = f"Job Description:\n{state['job_description']}\n\nGenerated Resume:\n{state['resume']}"
                review = resume_critic.run(review_prompt, temperature=0.2)
                review_data = json.loads(review)
            except:
                review_data = {"pass": True, "feedback": "JSON parse failed, assuming pass to proceed."}

            if review_data.get("pass", False) or step == MAX_ITERATIONS_RESUME - 1:
                state["logs"].append("✓ Resume approved!")
                break
            else:
                feedback = review_data.get('feedback', '')
                state["logs"].append(f"✗ Critic feedback: Needs improvement. Refining...")
                state["feedback"] = feedback

        # Formatting Phase
        state["current_phase"] = "Formatting Document"
        state["logs"].append("Generating Word Document (.docx)...")
        if state["resume"]:
            state["docx_path"] = create_resume_docx(state["resume"], job_id)
            
        state["status"] = "completed"
        state["current_phase"] = "Complete"
        state["logs"].append("Process finished successfully.")
        
    except Exception as e:
        state["status"] = "error"
        state["current_phase"] = "Error"
        state["logs"].append(f"Fatal Error: {str(e)}")


# ---------------------------------------------------
# Routes
# ---------------------------------------------------

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/resume/generate', methods=['POST'])
def generate_resume():
    target_role = request.form.get('target_role', '')
    job_description = request.form.get('job_description', '')
    
    # Process uploaded files to extract text
    files = request.files.getlist('files')
    extracted_text = []
    
    for file in files:
        if file.filename:
            text = extract_text_from_file(file)
            extracted_text.append(f"--- CONTENT FROM {file.filename} ---\n{text}")
            
    current_resume = "\n\n".join(extracted_text)
    
    if not job_description and not current_resume and not target_role:
        return jsonify({'error': 'Please provide at least a target role, job description, or upload a resume.'}), 400
    
    job_id = str(uuid.uuid4())
    
    # Start Agent Thread
    thread = threading.Thread(
        target=run_resume_agent, 
        args=(target_role, job_description, current_resume, job_id)
    )
    thread.daemon = True
    thread.start()
    
    return jsonify({'job_id': job_id})

@app.route('/api/resume/status/<job_id>')
def resume_status(job_id):
    if job_id in resume_jobs:
        job_data = resume_jobs[job_id]
        return jsonify({
            'status': job_data['status'],
            'current_phase': job_data['current_phase'],
            'current_iteration': job_data['current_iteration'],
            'logs': job_data['logs'],
            'has_file': job_data['docx_path'] is not None
        })
    return jsonify({'status': 'not_found'})

@app.route('/api/resume/download/<job_id>')
def download_resume(job_id):
    if job_id not in resume_jobs:
        return jsonify({'error': 'Job not found'}), 404
    
    job = resume_jobs[job_id]
    if not job.get('docx_path'):
        return jsonify({'error': 'Resume not ready'}), 404
    
    return send_file(job['docx_path'], as_attachment=True)

@app.route('/api/resume/text/<job_id>')
def resume_text(job_id):
    if job_id in resume_jobs:
        return jsonify({'resume': resume_jobs[job_id].get('resume', '')})
    return jsonify({'resume': ''})

if __name__ == '__main__':
    print("\n" + "="*50)
    print("📄 AI Resume Builder Started!")
    print("   Powered by Google Gemini GenAI")
    print("="*50 + "\n")
    app.run(debug=True, host='0.0.0.0', port=5000)